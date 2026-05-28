from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


BASE = Path(__file__).resolve().parent
DOCX = BASE / "车金_AI智能客服售前跟进系统_技术服务合同_2026-05-26.docx"
PDF = BASE / "车金_AI智能客服售前跟进系统_技术服务合同_2026-05-26.pdf"


pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def clean(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def body_elements(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            for p in doc.paragraphs:
                if p._p is child:
                    yield ("p", p)
                    break
        elif child.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._tbl is child:
                    yield ("tbl", table)
                    break


def make_styles():
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "CNBody",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=15,
        alignment=TA_LEFT,
        spaceAfter=5,
    )
    title = ParagraphStyle(
        "CNTitle",
        parent=base,
        fontSize=19,
        leading=26,
        alignment=TA_CENTER,
        spaceAfter=12,
        textColor=colors.HexColor("#0B2545"),
    )
    subtitle = ParagraphStyle(
        "CNSubtitle",
        parent=base,
        fontSize=10.5,
        leading=15,
        alignment=TA_CENTER,
        spaceAfter=14,
        textColor=colors.HexColor("#555555"),
    )
    h1 = ParagraphStyle(
        "CNH1",
        parent=base,
        fontSize=14,
        leading=20,
        spaceBefore=12,
        spaceAfter=7,
        textColor=colors.HexColor("#1F4D78"),
    )
    h2 = ParagraphStyle(
        "CNH2",
        parent=base,
        fontSize=12,
        leading=17,
        spaceBefore=9,
        spaceAfter=5,
        textColor=colors.HexColor("#2E74B5"),
    )
    return base, title, subtitle, h1, h2


def add_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("STSong-Light", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawCentredString(letter[0] / 2, 0.45 * inch, f"AI智能客服售前跟进系统技术服务合同  |  第 {doc.page} 页")
    canvas.restoreState()


def convert():
    source = Document(DOCX)
    base, title, subtitle, h1, h2 = make_styles()
    story = []
    seen_title = False

    for kind, obj in body_elements(source):
        if kind == "p":
            text = obj.text.strip()
            if not text:
                continue
            if text.startswith("附件") and story:
                story.append(PageBreak())
                story.append(Paragraph(clean(text), h1))
                continue
            if text.startswith("第") and "条 " in text:
                story.append(Paragraph(clean(text), h1))
            elif not seen_title and text.startswith("AI智能客服售前跟进系统技术服务合同"):
                story.append(Paragraph(clean(text), title))
                seen_title = True
            elif text.startswith("车金项目"):
                story.append(Paragraph(clean(text), subtitle))
            else:
                story.append(Paragraph(clean(text), base))
        else:
            data = []
            for row in obj.rows:
                data.append([Paragraph(clean(cell.text.strip()), base) for cell in row.cells])
            if data:
                table = Table(data, hAlign="CENTER", repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                            ("FONTSIZE", (0, 0), (-1, -1), 9.2),
                            ("LEADING", (0, 0), (-1, -1), 12),
                            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BFC7D1")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 5),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 8))

    pdf = SimpleDocTemplate(
        str(PDF),
        pagesize=letter,
        rightMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
        title="车金_AI智能客服售前跟进系统_技术服务合同",
        author="Codex",
    )
    pdf.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    print(PDF)


if __name__ == "__main__":
    convert()
